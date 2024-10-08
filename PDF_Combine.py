# import streamlit as st
# from PyPDF2 import PdfMerger
# from openpyxl import load_workbook
# from fpdf import FPDF
# from PIL import Image
# from io import BytesIO
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# import io
# from datetime import datetime
# import tempfile

# # Page configuration - should be at the top
# st.set_page_config(
#     page_title="Nicola's PDF Puzzle",
#     page_icon="📄",
#     layout="centered",
#     initial_sidebar_state="auto"
# )

# # Main title and subheader
# st.title("📄 Nicola's PDF Puzzle")
# st.subheader("From chaos to order—one PDF at a time! 🚀")

# # Helper functions
# def convert_word_to_pdf(docx_file):
#     try:
#         # Handle if the input file is BytesIO
#         if isinstance(docx_file, BytesIO):
#             # Save the BytesIO content to a temporary file
#             with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
#                 temp_docx.write(docx_file.getbuffer())
#                 temp_docx_path = temp_docx.name

#             # Read the temporary DOCX file
#             doc = Document(temp_docx_path)

#             # Check if the document is empty
#             if not doc.paragraphs:
#                 raise ValueError("The DOCX file is empty.")

#             # Create a PDF in memory
#             pdf_output = BytesIO()
#             c = canvas.Canvas(pdf_output, pagesize=letter)
#             width, height = letter

#             # Write each paragraph of the DOCX file to the PDF
#             text = c.beginText(40, height - 40)
#             text.setFont("Helvetica", 12)

#             for para in doc.paragraphs:
#                 text.textLine(para.text)

#             c.drawText(text)
#             c.showPage()
#             c.save()

#             pdf_output.seek(0)
#             return pdf_output

#         else:
#             raise ValueError("The input file is not a valid DOCX file.")
#     except Exception as e:
#         st.error(f"⚠️ An error occurred while converting DOCX to PDF: {str(e)}")
#         return None

# def convert_excel_to_pdf(excel_file):
#     try:
#         # Handle if the input file is BytesIO
#         if isinstance(excel_file, BytesIO):
#             # Load the workbook from the BytesIO object
#             wb = load_workbook(excel_file)
#             sheet = wb.active

#             # Create a PDF in memory
#             pdf_output = BytesIO()
#             c = canvas.Canvas(pdf_output, pagesize=letter)
#             width, height = letter

#             # Write each row of the Excel file to the PDF
#             text = c.beginText(40, height - 40)
#             text.setFont("Helvetica", 12)

#             for row in sheet.iter_rows(values_only=True):
#                 row_text = ' '.join([str(cell) for cell in row if cell is not None])
#                 text.textLine(row_text)

#             c.drawText(text)
#             c.showPage()
#             c.save()

#             pdf_output.seek(0)
#             return pdf_output

#         else:
#             raise ValueError("The input file is not a valid Excel file.")
#     except Exception as e:
#         st.error(f"⚠️ An error occurred while converting Excel to PDF: {str(e)}")
#         return None

# def convert_image_to_pdf(image_file):
#     try:
#         img = Image.open(image_file)
#         pdf_output = io.BytesIO()
#         img.convert('RGB').save(pdf_output, format='PDF')
#         pdf_output.seek(0)
#         return pdf_output if pdf_output.getbuffer().nbytes > 0 else None
#     except Exception as e:
#         st.error(f"⚠️ An error occurred while converting image to PDF: {str(e)}")
#         return None

# # Instructions
# st.write("""
# **Upload up to 15 Images, Word, Excel, or PDF documents below.**
# We'll help you combine them into one single, neat PDF file that you can download.
# """)

# # File upload
# uploaded_files = st.file_uploader("Choose files", type=["pdf", "png", "jpg", "jpeg", "docx", "xlsx"], accept_multiple_files=True)

# if uploaded_files:
#     # Display a limit on the number of files
#     if len(uploaded_files) > 15:
#         st.error("🚨 You can upload a maximum of 15 files. Please try again.")
#     else:
#         st.write("All files uploaded successfully! Press the **Create PDF** button below to merge them.")

#         # Display files with dropdowns for reordering
#         file_names = [file.name for file in uploaded_files]
#         reordered_files = []

#         st.write("### Files Uploaded:")
#         for idx, file_name in enumerate(file_names):
#             selected_position = st.selectbox(
#                 label=f"Position for Attachment {idx + 1}: {file_name}",
#                 options=list(range(1, len(file_names) + 1)),
#                 index=idx,
#                 key=f"select_{idx}"
#             )
#             reordered_files.append((selected_position, uploaded_files[idx]))

#         # Sort files based on selected positions
#         reordered_files.sort(key=lambda x: x[0])
#         sorted_files = [file[1] for file in reordered_files]

#         # Button to start the merging process
#         if st.button("🎉 Create PDF!"):
#             merger = PdfMerger()
#             for file in sorted_files:
#                 file_type = file.name.split('.')[-1].lower()
#                 pdf_file = None
                
#                 if file_type == 'pdf':
#                     pdf_file = file
#                 elif file_type == 'docx':
#                     pdf_file = convert_word_to_pdf(file)
#                 elif file_type == 'xlsx':
#                     pdf_file = convert_excel_to_pdf(file)
#                 elif file_type in ['png', 'jpg', 'jpeg']:
#                     pdf_file = convert_image_to_pdf(file)
                
#                 if pdf_file:
#                     try:
#                         merger.append(pdf_file)
#                     except Exception as e:
#                         st.error(f"Error merging {file.name}: {str(e)}")
#                 else:
#                     st.error(f"Conversion failed for {file.name}. Skipping this file.")

#             merged_pdf = io.BytesIO()
#             merger.write(merged_pdf)
#             merged_pdf.seek(0)
#             merger.close()

#             # Generate timestamp for file name
#             timestamp = datetime.now().strftime("%Y%m%d%H%M")
#             file_name = f"merged_document_{timestamp}.pdf"

#             st.success("🎉 PDF created successfully! Download your merged PDF below.")
#             st.download_button(
#                 label="📥 Download Merged PDF",
#                 data=merged_pdf,
#                 file_name=file_name,
#                 mime="application/pdf"
#             )


import streamlit as st
import io
from zipfile import ZipFile
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO, BytesIO
import base64
import pdf2image
import pytesseract
from pytesseract import Output, TesseractError
from PyPDF2 import PdfMerger, PdfWriter, PdfReader
from openpyxl import load_workbook
from fpdf import FPDF
from PIL import Image
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from datetime import datetime
import tempfile

# Page configuration - should be at the top
st.set_page_config(
    page_title="Nicola's PDF Puzzle",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="auto"
)

# Main title and subheader
st.title("📄 Nicola's PDF Puzzle")
st.subheader("From chaos to order—one PDF at a time! 🚀")

# Helper functions
def convert_word_to_pdf(docx_file):
    try:
        if isinstance(docx_file, BytesIO):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                temp_docx.write(docx_file.getbuffer())
                temp_docx_path = temp_docx.name
            doc = Document(temp_docx_path)
            if not doc.paragraphs:
                raise ValueError("The DOCX file is empty.")
            pdf_output = BytesIO()
            c = canvas.Canvas(pdf_output, pagesize=letter)
            width, height = letter
            text = c.beginText(40, height - 40)
            text.setFont("Helvetica", 12)
            for para in doc.paragraphs:
                text.textLine(para.text)
            c.drawText(text)
            c.showPage()
            c.save()
            pdf_output.seek(0)
            return pdf_output
        else:
            raise ValueError("The input file is not a valid DOCX file.")
    except Exception as e:
        st.error(f"⚠️ An error occurred while converting DOCX to PDF: {str(e)}")
        return None

def convert_excel_to_pdf(excel_file):
    try:
        if isinstance(excel_file, BytesIO):
            wb = load_workbook(excel_file)
            sheet = wb.active
            pdf_output = BytesIO()
            c = canvas.Canvas(pdf_output, pagesize=letter)
            width, height = letter
            text = c.beginText(40, height - 40)
            text.setFont("Helvetica", 12)
            for row in sheet.iter_rows(values_only=True):
                row_text = ' '.join([str(cell) for cell in row if cell is not None])
                text.textLine(row_text)
            c.drawText(text)
            c.showPage()
            c.save()
            pdf_output.seek(0)
            return pdf_output
        else:
            raise ValueError("The input file is not a valid Excel file.")
    except Exception as e:
        st.error(f"⚠️ An error occurred while converting Excel to PDF: {str(e)}")
        return None

def convert_image_to_pdf(image_file):
    try:
        img = Image.open(image_file)
        pdf_output = BytesIO()
        img.convert('RGB').save(pdf_output, format='PDF')
        pdf_output.seek(0)
        return pdf_output if pdf_output.getbuffer().nbytes > 0 else None
    except Exception as e:
        st.error(f"⚠️ An error occurred while converting image to PDF: {str(e)}")
        return None

@st.cache_data
def images_to_txt(path, language):
    images = pdf2image.convert_from_bytes(path)
    all_text = []
    for i in images:
        pil_im = i
        text = pytesseract.image_to_string(pil_im, lang=language)
        all_text.append(text)
    return all_text, len(all_text)

def add_ocr_text_to_pdf(pdf_bytes, ocr_texts, ocr_boxes):
    pdf_reader = PdfReader(BytesIO(pdf_bytes))
    pdf_writer = PdfWriter()

    for page_num, page in enumerate(pdf_reader.pages):
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        can.setFont("Helvetica", 12)
        can.setFillAlpha(0.0)  # Set transparency to make text invisible
        for box in ocr_boxes[page_num]:
            x, y, w, h, text = box
            can.drawString(x, y, text)
        can.save()

        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        page.merge_page(overlay_pdf.pages[0])
        pdf_writer.add_page(page)

    output_pdf = BytesIO()
    pdf_writer.write(output_pdf)
    output_pdf.seek(0)
    return output_pdf

def extract_ocr_boxes(images, language):
    ocr_boxes = []
    for img in images:
        data = pytesseract.image_to_data(img, lang=language, output_type=Output.DICT)
        boxes = []
        for i in range(len(data['text'])):
            if int(data['conf'][i]) > 60:  # Confidence threshold
                (x, y, w, h) = (data['left'][i], data['top'][i], data['width'][i], data['height'][i])
                text = data['text'][i]
                boxes.append((x, y, w, h, text))
        ocr_boxes.append(boxes)
    return ocr_boxes

# Instructions
st.write("""
**Upload up to 15 Images, Word, Excel, or PDF documents below.**
We'll help you combine them into one single, neat PDF file that you can download.
""")

# File upload
uploaded_files = st.file_uploader("Choose files", type=["pdf", "png", "jpg", "jpeg", "docx", "xlsx"], accept_multiple_files=True)

if uploaded_files:
    if len(uploaded_files) > 15:
        st.error("🚨 You can upload a maximum of 15 files. Please try again.")
    else:
        st.write("All files uploaded successfully! Press the **Create PDF** button below to merge them.")
        file_names = [file.name for file in uploaded_files]
        reordered_files = []
        st.write("### Files Uploaded:")
        for idx, file_name in enumerate(file_names):
            selected_position = st.selectbox(
                label=f"Position for Attachment {idx + 1}: {file_name}",
                options=list(range(1, len(file_names) + 1)),
                index=idx,
                key=f"select_{idx}"
            )
            reordered_files.append((selected_position, uploaded_files[idx]))
        reordered_files.sort(key=lambda x: x[0])
        sorted_files = [file[1] for file in reordered_files]

        if st.button("🎉 Create PDF!"):
            merger = PdfMerger()
            for file in sorted_files:
                file_type = file.name.split('.')[-1].lower()
                pdf_file = None
                if file_type == 'pdf':
                    pdf_file = file
                elif file_type == 'docx':
                    pdf_file = convert_word_to_pdf(file)
                elif file_type == 'xlsx':
                    pdf_file = convert_excel_to_pdf(file)
                elif file_type in ['png', 'jpg', 'jpeg']:
                    pdf_file = convert_image_to_pdf(file)
                if pdf_file:
                    try:
                        merger.append(pdf_file)
                    except Exception as e:
                        st.error(f"Error merging {file.name}: {str(e)}")
                else:
                    st.error(f"Conversion failed for {file.name}. Skipping this file.")
            merged_pdf = BytesIO()
            merger.write(merged_pdf)
            merged_pdf.seek(0)
            merger.close()

            # Add OCR to the merged PDF
            images = pdf2image.convert_from_bytes(merged_pdf.getvalue())
            ocr_texts, num_pages = images_to_txt(merged_pdf.getvalue(), 'eng')
            ocr_boxes = extract_ocr_boxes(images, 'eng')
            if ocr_texts:
                merged_pdf_with_ocr = add_ocr_text_to_pdf(merged_pdf.getvalue(), ocr_texts, ocr_boxes)
                timestamp = datetime.now().strftime("%Y%m%d%H%M")
                file_name = f"merged_document_{timestamp}.pdf"
                st.success("🎉 PDF created successfully with OCR! Download your merged PDF below.")
                st.download_button(
                    label="📥 Download Merged PDF with OCR",
                    data=merged_pdf_with_ocr,
                    file_name=file_name,
                    mime="application/pdf"
                )
